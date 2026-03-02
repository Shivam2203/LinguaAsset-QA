"""
Document Processor - Extracts and processes questions from Word documents
"""

import re
from docx import Document
from typing import List, Dict, Any, Optional

def extract_questions_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    Extract questions from a Word document
    
    Args:
        docx_path: Path to Word document
        
    Returns:
        List of questions with metadata
    """
    questions = []
    
    try:
        doc = Document(docx_path)
        
        # Extract from tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if _is_question(text) or _is_blank(text):
                        questions.append({
                            'type': 'table',
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'col_idx': col_idx,
                            'text': text,
                            'field_name': _extract_field_name(text)
                        })
        
        # Extract from paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text and (_is_question(text) or _is_blank(text)):
                questions.append({
                    'type': 'paragraph',
                    'para_idx': para_idx,
                    'text': text,
                    'field_name': _extract_field_name(text)
                })
    
    except Exception as e:
        print(f"Error extracting questions: {e}")
    
    return questions

def _is_question(text: str) -> bool:
    """Check if text is a question"""
    question_patterns = [
        r'.*?\?',           # Ends with question mark
        r'^what.*',         # Starts with what
        r'^who.*',          # Starts with who
        r'^when.*',         # Starts with when
        r'^where.*',        # Starts with where
        r'^how.*',          # Starts with how
        r'^why.*',          # Starts with why
        r'please.*?:',      # Please with colon
        r'^\d+\..*',        # Numbered list
    ]
    
    text_lower = text.lower()
    for pattern in question_patterns:
        if re.match(pattern, text_lower):
            return True
    return False

def _is_blank(text: str) -> bool:
    """Check if text contains a blank to fill"""
    blank_patterns = [
        r'_{3,}',           # underscores
        r'\[.*\]',          # brackets
        r'\(.*\)',          # parentheses
        r'<.*>',            # angle brackets
    ]
    
    for pattern in blank_patterns:
        if re.search(pattern, text):
            return True
    return False

def _extract_field_name(text: str) -> str:
    """Extract field name from question text"""
    # Remove common prefixes
    text = re.sub(r'^\d+\.\s*', '', text)
    text = re.sub(r'^please\s+', '', text, flags=re.IGNORECASE)
    
    # Take first part before colon or question mark
    if ':' in text:
        return text.split(':')[0].strip()
    if '?' in text:
        return text.split('?')[0].strip()
    
    return text[:50]  # First 50 chars