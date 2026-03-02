"""
Word Styler - Applies professional formatting to Word documents
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Optional, Dict

class WordStyler:
    """Apply professional formatting to Word documents"""
    
    def format_document(self, input_path: str, output_path: Optional[str] = None) -> Dict:
        """
        Format a Word document with professional styling
        
        Args:
            input_path: Path to input document
            output_path: Path to output document (if None, overwrite input)
            
        Returns:
            Dict with formatting results
        """
        if output_path is None:
            output_path = input_path
        
        result = {
            'success': False,
            'paragraphs_formatted': 0,
            'tables_formatted': 0
        }
        
        try:
            # Open document
            doc = Document(input_path)
            
            # Set page margins
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.0)
                section.right_margin = Cm(2.5)
            
            # Format paragraphs
            font_name = 'FangSong'  # Chinese font
            for para in doc.paragraphs:
                if para.text.strip():
                    # Set paragraph formatting
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.first_line_indent = Cm(0.75)
                    
                    # Set font for all runs
                    for run in para.runs:
                        run.font.name = font_name
                        run.font.size = Pt(12)
                        # Set East Asian font
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    
                    result['paragraphs_formatted'] += 1
            
            # Format tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = font_name
                                run.font.size = Pt(11)
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                result['tables_formatted'] += 1
            
            # Save document
            doc.save(output_path)
            result['success'] = True
            
        except Exception as e:
            print(f"Error formatting document: {e}")
            result['error'] = str(e)
        
        return result

# Global instance
_styler = None

def get_styler() -> WordStyler:
    """Get styler instance"""
    global _styler
    if _styler is None:
        _styler = WordStyler()
    return _styler