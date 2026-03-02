"""
Table Formatter - Creates and formats tables in Word documents
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any, Optional

class TableFormatter:
    """Create and format tables in Word documents"""
    
    def create_table(self, doc: Document, table_data: Dict[str, Any], 
                     title: str = "") -> Optional[Document]:
        """
        Create a formatted table in Word document
        
        Args:
            doc: Word document object
            table_data: Dict with 'headers' and 'rows'
            title: Optional table title
            
        Returns:
            Updated document
        """
        try:
            headers = table_data.get('headers', [])
            rows = table_data.get('rows', [])
            
            if not headers or not rows:
                return doc
            
            # Add title if provided
            if title:
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.bold = True
                run.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Create table
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set column widths
            for col in table.columns:
                col.width = Inches(1.5)
            
            # Add headers
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = str(header)
                # Format header
                for para in header_row.cells[i].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            
            # Add data rows
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(headers):
                        cell = table.rows[row_idx + 1].cells[col_idx]
                        cell.text = str(cell_data) if cell_data else ''
                        # Format cell
                        for para in cell.paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in para.runs:
                                run.font.size = Pt(11)
            
            return doc
            
        except Exception as e:
            print(f"Error creating table: {e}")
            return doc
    
    def should_display_table(self, field_name: str, context: str = "") -> bool:
        """Determine if field should be displayed as table"""
        table_fields = [
            'shareholder', 'fund', 'product', 'team', 'list',
            '股东', '基金', '产品', '团队', '列表'
        ]
        
        field_lower = field_name.lower()
        for field in table_fields:
            if field in field_lower:
                return True
        
        return False

# Global instance
_formatter = None

def get_table_formatter() -> TableFormatter:
    """Get table formatter instance"""
    global _formatter
    if _formatter is None:
        _formatter = TableFormatter()
    return _formatter